from io import BytesIO
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

class UseCaseDocGenerator:
    def __init__(self, data, table_width_cm=17.8):
        self.data = data
        self.table_width_cm = table_width_cm

    # -------------------------
    # Helper functies
    # -------------------------
    @staticmethod
    def set_cell_bg(cell, color):
        if color:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color)
            tcPr.append(shd)

    @staticmethod
    def set_table_borders(table, border_color="000000"):
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = tblPr.find(qn('w:tblBorders'))
        if tblBorders is None:
            tblBorders = OxmlElement('w:tblBorders')
            tblPr.append(tblBorders)
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = tblBorders.find(qn(f'w:{border_name}'))
            if border is None:
                border = OxmlElement(f'w:{border_name}')
                tblBorders.append(border)
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)

    # -------------------------
    # Flexibele tabel functie
    # -------------------------
    def add_flexible_table(self, doc, title, table_data):
        if not table_data:
            return None

        if title:
            doc.add_paragraph(title, style="Heading 2")

        max_cols = max(len(row) for row in table_data)
        table = doc.add_table(rows=len(table_data), cols=max_cols)

        # Stel hele tabelbreedte in via XML
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(int(self.table_width_cm * 567)))
        tblW.set(qn('w:type'), 'dxa')

        for i, row_data in enumerate(table_data):
            row_len = len(row_data)
            for j, cell_data in enumerate(row_data):
                cell = table.cell(i, j)
                cell.text = cell_data.get("text", "")
                self.set_cell_bg(cell, cell_data.get("bg_color"))

                # Breedte per cel
                width_pct = cell_data.get("width")
                if width_pct:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcW = OxmlElement('w:tcW')
                    tcW.set(qn('w:w'), str(int(self.table_width_cm * width_pct * 567)))
                    tcW.set(qn('w:type'), 'dxa')
                    tcPr.append(tcW)

            # Merge lege cellen naar laatste cel
            if row_len < max_cols:
                last_cell = table.cell(i, row_len - 1)
                for k in range(row_len, max_cols):
                    last_cell.merge(table.cell(i, k))

        self.set_table_borders(table)
        return table

    # -------------------------
    # Document genereren
    # -------------------------
    def generate_docx_bytes(self):
        doc = Document()
        for section in doc.sections:
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Metadata
        self.add_flexible_table(doc, "", self.data.get("metadata", []))
        doc.add_paragraph("")
        # Preconditions
        self.add_flexible_table(doc, "", self.data.get("preconditions", []))
        doc.add_paragraph("")
        # Basic Flow
        self.add_flexible_table(doc, "", self.data.get("basic_flow", []))
        doc.add_paragraph("")
        # Alternate Flow
        self.add_flexible_table(doc, "", self.data.get("alternate_flows", []))
        doc.add_paragraph("")
        # Exception Flow
        self.add_flexible_table(doc, "", self.data.get("exception_flows", []))
        doc.add_paragraph("")
        # Postconditions
        self.add_flexible_table(doc, "", self.data.get("postconditions", []))

        # Opslaan in bytes
        byte_io = BytesIO()
        doc.save(byte_io)
        byte_io.seek(0)
        return byte_io.getvalue()
