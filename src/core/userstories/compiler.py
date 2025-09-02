import io
from docx import Document


class UserStoryCompiler:
    def __init__(self, user_stories: list[dict]):
        self.user_stories = user_stories

    def to_string(self) -> str:
        """Retourneert de user stories als platte tekst (string)."""
        lines = []
        for us in self.user_stories:
            lines.append(f"ID: {us['id']}")
            lines.append(f"Titel: {us['title']}")
            lines.append(f"Als een {us['user_story']['as_a']}, "
                         f"wil ik {us['user_story']['i_want']} "
                         f"zodat {us['user_story']['so_that']}.")
            lines.append(f"Beschrijving: {us['description']}")
            lines.append("Acceptatiecriteria:")
            for crit in us["acceptance_criteria"]:
                lines.append(f"  - {crit}")
            lines.append("")  # lege regel
        return "\n".join(lines)

    def to_txt(self) -> bytes:
        """Retourneert de user stories als bytes (txt)."""
        text = self.to_string()
        return text.encode("utf-8")

    def to_docx(self) -> bytes:
        """Retourneert de user stories als bytes (docx)."""
        doc = Document()
        for us in self.user_stories:
            doc.add_heading(us["title"], level=1)
            doc.add_paragraph(f"ID: {us['id']}")
            doc.add_paragraph(f"Als een {us['user_story']['as_a']}, "
                              f"wil ik {us['user_story']['i_want']} "
                              f"zodat {us['user_story']['so_that']}.")
            doc.add_paragraph(f"Beschrijving: {us['description']}")
            doc.add_heading("Acceptatiecriteria", level=2)
            for crit in us["acceptance_criteria"]:
                doc.add_paragraph(crit, style="List Bullet")
            doc.add_paragraph("")  # lege regel

        # Schrijf in geheugen (ipv bestand)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
