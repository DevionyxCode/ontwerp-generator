from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm

# -------------------------
# Helper functies
# -------------------------
def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p


def add_list(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def set_table_borders(table, border_color="000000"):
    """Zet alle randen van een tabel aan"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is None:
        tblBorders = OxmlElement('w:tblBorders')
        tblPr.append(tblBorders)

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = tblBorders.find(qn(f'w:{border_name}'))
        if border is None:
            border = OxmlElement(f'w:{border_name}')
            tblBorders.append(border)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), border_color)


def set_cell_bg(cell, color):
    """Stel achtergrondkleur in voor één cel"""
    if color:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)  # HEX kleur
        tcPr.append(shd)


# -------------------------
# Flexibele tabel functie
# -------------------------
def add_flexible_table(doc, title, table_data, total_width_cm=17.8):
    """
    Maak een tabel volledig flexibel per cel met achtergrondkleur en breedte per cel.
    Merge cells automatisch als er minder cellen zijn dan de max kolommen.
    De tabel wordt breder (default 18 cm, bijna volledige A4 breedte).
    """
    if not table_data:
        return None

    if title:
        doc.add_paragraph(title, style="Heading 2")

    # Max aantal kolommen
    max_cols = max(len(row) for row in table_data)

    table = doc.add_table(rows=len(table_data), cols=max_cols)

    # Stel hele tabelbreedte in via XML
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(int(total_width_cm * 567)))  # 1 cm ≈ 567 Twips
    tblW.set(qn('w:type'), 'dxa')

    for i, row_data in enumerate(table_data):
        row_len = len(row_data)
        for j, cell_data in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = cell_data.get("text", "")
            set_cell_bg(cell, cell_data.get("bg_color"))

            # Breedte per cel
            width_pct = cell_data.get("width")
            if width_pct:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(total_width_cm * width_pct * 567)))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

        # Merge lege cellen naar laatste cel
        if row_len < max_cols:
            last_cell = table.cell(i, row_len - 1)
            for k in range(row_len, max_cols):
                last_cell.merge(table.cell(i, k))

    set_table_borders(table)
    return table

# -------------------------
# Build document
# -------------------------
def build_usecase_doc(data, output_path="usecase.docx"):
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Titel en meta
    add_flexible_table(doc, "", data.get("metadata", []))
    doc.add_paragraph("")
    add_flexible_table(doc, "", data.get("preconditions", []))

    # Basic Flow
    doc.add_paragraph("")
    add_flexible_table(doc, "", data.get("basic_flow", []))

    # Alternate Flow
    doc.add_paragraph("")
    add_flexible_table(doc, "", data.get("alternate_flows", []))

    # Exception Flow
    doc.add_paragraph("")
    add_flexible_table(doc, "", data.get("exception_flows", []))

    # Postconditions
    doc.add_paragraph("")
    add_flexible_table(doc, "", data.get("postconditions", []))

    doc.save(output_path)
    print(f"Document saved to {output_path}")

#        {
#        "title": "Gebruiker blokkeren",
#        "version": "0.1",
#        "end_goal": "De toegang van een gebruiker tot het systeem tijdelijk beperken.",
#        "description": "Een beheerder blokkeert een actieve gebruiker ...",
#        "actors": ["Beheerder"],
#    },
# -------------------------
# Voorbeeld JSON
# -------------------------
sample_json = {
    "metadata": [
        [
            {"text": "Use Case:", "bg_color": "CCCCCC", "width": 0.25},
            {"text": "Gebruiker blokkeren", "width": 0.4},
            {"text": "Version No:", "bg_color": "CCCCCC", "width": 0.15},
            {"text": "0.1", "width": 0.25}
        ],
        [
            {"text": "End goal:", "bg_color": "CCCCCC", "width": 0.25},
            {"text": "De toegang van een gebruiker tot het systeem tijdelijk beperken.", "width": 0.75},
        ],
        [
            {"text": "Description:", "bg_color": "CCCCCC", "width": 0.25},
            {"text": "Een beheerder blokkeert een actieve gebruiker, waardoor deze niet meer kan inloggen of acties uitvoeren in het systeem totdat deze weer wordt gedeblokkeerd.", "width": 0.75},
        ],
        [
            {"text": "Actors:", "bg_color": "CCCCCC", "width": 0.25},
            {"text": "Beheerder", "width": 0.75},
        ]
    ],
    "preconditions": [
        [
            {"text": "Preconditions", "bg_color": "CCCCCC"},
        ],
        [
            {"text": '''  1.  De beheerder is succesvol ingelogd in het systeem.\n  2.  De beheerder heeft de juiste autorisatie om gebruikers te blokkeren.\n  3.  De te blokkeren gebruiker bestaat in het systeem.\n  4.  De te blokkeren gebruiker heeft op dit moment een actieve status.\n  5.  De te blokkeren gebruiker is niet de huidige ingelogde beheerder.
            '''},
        ]
    ],
    "basic_flow": [
        [
            {"text": '''Basic Flow \nThe optimal or normal ("good day") flow of events.\nThe basic flow of events should describe the events that walk through a successful scenario.\nThe basic flow should not include “and/if scenarios”''', "bg_color": "CCCCCC"}
        ],
        [
            {"text": "Step", "bg_color": "CCCCCC", "width": 0.15},
            {"text": "User Actions", "bg_color": "CCCCCC", "width": 0.30},
            {"text": "System Actions", "bg_color": "CCCCCC", "width": 0.55}
        ],
        [
            {"text": "1", "width": 0.15},
            {"text": "Beheerder navigeert naar overzicht", "width": 0.30},
            {"text": "Systeem toont gebruikers", "width": 0.55}
        ],
        [
            {"text": "2", "width": 0.15},
            {"text": "Beheerder selecteert gebruiker", "width": 0.30},
            {"text": "Systeem toont detailpagina", "width": 0.55}
        ]
    ],
    "alternate_flows": [
        [
            {
                "text": '''Alternate Flow \nThere may be more than one.''',
                "bg_color": "CCCCCC"}
        ],
        [
            {"text": "Step", "bg_color": "CCCCCC", "width": 0.15},
            {"text": "User Actions", "bg_color": "CCCCCC", "width": 0.30},
            {"text": "System Actions", "bg_color": "CCCCCC", "width": 0.55}
        ],
        [
            {"text": "1", "width": 0.15},
            {"text": "Meerdere gebruikers selecteren", "width": 0.30},
            {"text": "Systeem highlight", "width": 0.55}
        ]
    ],
    "exception_flows": [
        [
            {
                "text": '''Exception Flow  \nThere may be more than one.\nIdentify system and data error conditions that could occur for each step in the normal and alternate flow''',
                "bg_color": "CCCCCC"}
        ],
        [
            {"text": "1", "width": 0.15},
            {"text": "Beheerder probeert zichzelf te blokkeren", "width": 0.30},
            {"text": "Systeem toont foutmelding", "width": 0.55}
        ]
    ],
    "postconditions": [
        [
            {"text": "Preconditions", "bg_color": "CCCCCC"},
        ],
        [
            {"text": ''' 1.	De gebruiker(s) is/zijn geblokkeerd en kunnen niet inloggen.\n2.	Eventuele actieve sessies van de geblokkeerde gebruiker(s) zijn beëindigd.\n3.	De status van de gebruiker(s) is in het systeem gewijzigd naar "geblokkeerd".\n4.	Er is een auditlog aangemaakt van de blokkering(sactie).\n5.	De beheerder ontvangt een bevestiging van de succesvolle blokkering.'''},
        ]
    ],
}

# -------------------------
# Main
# -------------------------
if __name__ == "__main__":
    build_usecase_doc(sample_json, "output_usecase.docx")
